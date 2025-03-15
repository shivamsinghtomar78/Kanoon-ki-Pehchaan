import streamlit as st
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import SystemMessage, HumanMessage, AIMessage
from dotenv import load_dotenv
import os
import time
from datetime import datetime
import PyPDF2
import docx
from fpdf import FPDF

# Set the page title for the browser tab
st.set_page_config(page_title=" Kanoon ki Pehchaan", page_icon="⚖️")

# Load environment variables
load_dotenv()

# Initialize the Google Generative AI model
api_key = os.getenv("GOOGLE_API_KEY")
llm = ChatGoogleGenerativeAI(model="gemini-1.5-pro", google_api_key=api_key)

# Function to extract text from PDF
def extract_text_from_pdf(pdf_file):
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

# Function to analyze text using Google Generative AI
def analyze_text(text):
    response = llm([SystemMessage(content="Analyze the following text and provide a summary and key insights:"), HumanMessage(content=text)])
    return response.content

def create_pdf_report(analysis):
    pdf = FPDF()
    pdf.add_page()
    
    # Set font
    pdf.set_font('Helvetica', '', 12)
    
    # Add report title and generation date
    pdf.cell(200, 10, txt="PDF Analysis Report", ln=True, align='C')
    pdf.cell(200, 10, txt="Generated on: " + datetime.now().strftime("%Y-%m-%d %H:%M:%S"), ln=True, align='C')
    
    # Clean text by replacing problematic Unicode characters
    def clean_text_for_latin1(input_text):
        if not isinstance(input_text, str):
            return str(input_text)
            
        # Replace common Unicode characters with Latin-1 equivalents
        replacements = {
            '\u2013': '-',    # en dash
            '\u2014': '--',   # em dash
            '\u201c': '"',    # left double quotation mark
            '\u201d': '"',    # right double quotation mark
            '\u2018': "'",    # left single quotation mark
            '\u2019': "'",    # right single quotation mark
            '\u2026': '...',  # ellipsis
            # Add more replacements as needed
        }
        
        for unicode_char, replacement in replacements.items():
            input_text = input_text.replace(unicode_char, replacement)
            
        # Final fallback: replace any remaining non-Latin-1 characters with spaces
        return ''.join(c if ord(c) < 256 else ' ' for c in input_text)
    
    clean_analysis_content = clean_text_for_latin1(analysis)
    
    # Add analysis
    pdf.multi_cell(0, 10, txt="Analysis:\n" + clean_analysis_content)
    
    # Save the PDF
    pdf.output("analysis_report.pdf")

# Function to create a Word report
def create_word_report(analysis):
    doc = docx.Document()
    doc.add_heading('PDF Analysis Report', 0)
    doc.add_paragraph('Generated on: ' + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    doc.add_heading('Analysis', level=1)
    doc.add_paragraph(analysis)
    doc.save("analysis_report.docx")

# Custom CSS function
def local_css():
    st.markdown("""
    <style>
        /* Main container styling */
        .main {
            background: linear-gradient(135deg, #1e1e2f 0%, #2a2a40 100%);
            padding-bottom: 100px;
            color: #ffffff;
        }

        /* Header styling */
        .main-header {
            text-align: center;
            padding: 2rem 0;
            background: rgba(255, 255, 255, 0.1);
            border-radius: 15px;
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
            margin-bottom: 2rem;
            backdrop-filter: blur(10px);
            animation: fadeIn 1s ease-in-out;
        }

        .flag-stripe {
            height: 8px;
            background: linear-gradient(90deg, #ff9933 33%, white 33% 66%, #138808 66%);
            margin-bottom: 1rem;
            animation: slideIn 1s ease-in-out;
        }

        /* File uploader styling */
        .stFileUploader {
            background: rgba(255, 255, 255, 0.1);
            border-radius: 15px;
            padding: 1rem;
            border: 1px solid rgba(255, 255, 255, 0.2);
        }

        /* Button styling (applies to both st.button and st.download_button) */
        .stButton button, .stDownloadButton button {
            border-radius: 25px;
            padding: 0.5rem 1rem;
            background: linear-gradient(135deg, #ff9933 0%, #138808 100%);
            color: #ffffff;
            border: none;
            transition: all 0.3s ease;
        }

        .stButton button:hover, .stDownloadButton button:hover {
            transform: translateY(-2px);
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.2);
        }

        /* Analysis results styling */
        .analysis-results {
            background: rgba(248, 249, 250, 0.1);
            padding: 1rem;
            border-radius: 15px;
            border: 1px solid rgba(233, 236, 239, 0.2);
            color: #ffffff;
        }

        /* Ensure headers are white */
        h1, h2, h3, h4, h5, h6 {
            color: #ffffff;
        }

        /* Footer styling */
        .footer {
            position: fixed;
            bottom: 0;
            left: 0;
            right: 0;
            background: rgba(30, 30, 47, 0.9);
            padding: 1rem;
            text-align: center;
            border-top: 1px solid rgba(255, 255, 255, 0.1);
            font-size: 0.9em;
            color: #ffffff;
            backdrop-filter: blur(10px);
            animation: fadeIn 2s ease-in-out;
        }

        /* Animations */
        @keyframes fadeIn {
            from { opacity: 0; }
            to { opacity: 1; }
        }

        @keyframes slideIn {
            from { transform: translateX(-100%); }
            to { transform: translateX(0); }
        }
    </style>
    """, unsafe_allow_html=True)

# Streamlit UI
# Apply custom CSS
local_css()

# Initialize session state variables
if "current_file" not in st.session_state:
    st.session_state.current_file = None
if "pdf_summary" not in st.session_state:
    st.session_state.pdf_summary = None
if "analysis_time" not in st.session_state:
    st.session_state.analysis_time = 0

# Main content wrapper
st.markdown('<div class="main">', unsafe_allow_html=True)

# Custom header
st.markdown("""
<div class="main-header">
    <div class="flag-stripe"></div>
    <h1>⚖️ Kanoon ki Pehchaan</h1>
    <p>Your AI-powered Legal Document Analyzer</p>
</div>
""", unsafe_allow_html=True)

# Brief instruction
st.markdown("""
<p style="text-align: center; color: #ffffff;">
    Upload a PDF file to get AI-powered analysis and download the report.
</p>
""", unsafe_allow_html=True)

uploaded_file = st.file_uploader("Upload a PDF file", type="pdf")

if uploaded_file is not None:
    # Check if the file has changed
    if st.session_state.current_file != uploaded_file.name:
        st.session_state.current_file = uploaded_file.name
        st.session_state.pdf_summary = None
    
    # Extract text from PDF
    text = extract_text_from_pdf(uploaded_file)
    
    # Analyze text with timing
    if st.button("Analyze Text"):
        start_time = time.time()
        with st.spinner("Analyzing..."):
            analysis = analyze_text(text)
            st.session_state.pdf_summary = analysis
        end_time = time.time()
        st.session_state.analysis_time = end_time - start_time
        
        # Display analysis results
        st.markdown('<div class="analysis-results">', unsafe_allow_html=True)
        st.subheader("Analysis Results")
        st.write(analysis)
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Save analysis to PDF and Word
        create_pdf_report(analysis)
        create_word_report(analysis)
        
        # Provide download links
        st.download_button(
            label="Download PDF Report",
            data=open("analysis_report.pdf", "rb").read(),
            file_name="analysis_report.pdf",
            mime="application/pdf"
        )
        st.download_button(
            label="Download Word Report",
            data=open("analysis_report.docx", "rb").read(),
            file_name="analysis_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# Close main wrapper
st.markdown('</div>', unsafe_allow_html=True)

# Footer
st.markdown(f'''
<div class="footer">
    Analysis Time: {st.session_state.analysis_time:.1f}s | Powered by Google Generative AI
</div>
''', unsafe_allow_html=True)