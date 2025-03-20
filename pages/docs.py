import streamlit as st
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import SystemMessage, HumanMessage
from langchain.output_parsers import PydanticOutputParser
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from pydantic import BaseModel, Field
from typing import List
from dotenv import load_dotenv
import os
import time
from datetime import datetime
import PyPDF2
from fpdf import FPDF
from docx import Document
import io

# Load environment variables
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")

# Initialize the language model
llm = ChatGoogleGenerativeAI(model="gemini-1.5-pro", google_api_key=api_key)

# Step 1: Define Pydantic Models for Structured Output
class KeyPoint(BaseModel):
    point: str = Field(description="A key point extracted from the document.")

class Summary(BaseModel):
    summary: str = Field(description="A brief summary of the document content.")

class DocumentAnalysis(BaseModel):
    key_points: List[KeyPoint] = Field(description="List of key points from the document.")
    summary: Summary = Field(description="Summary of the document.")

# Step 2: Set Up LangChain with Output Parsing
parser = PydanticOutputParser(pydantic_object=DocumentAnalysis)

# Define the prompt with instructions for the LLM
prompt_template = """
Analyze the following text and extract key points and a summary.
{format_instructions}
Text: {text}
"""
prompt = PromptTemplate(
    template=prompt_template,
    input_variables=["text"],
    partial_variables={"format_instructions": parser.get_format_instructions()}
)

# Step 3: Chain the Components Together
chain = LLMChain(llm=llm, prompt=prompt, output_parser=parser)

# Function to analyze text and return structured data
def analyze_text_structured(text):
    output = chain.run(text=text)
    return output

# Function to extract text from a PDF file
def extract_text_from_pdf(pdf_file):
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

# Function to convert JSON to text format
def json_to_text(analysis):
    """
    Convert the structured JSON analysis into a simple text format.
    """
    text_output = ""

    # Add summary
    text_output += "=== Summary ===\n"
    text_output += f"{analysis.summary.summary}\n\n"

    # Add key points
    text_output += "=== Key Points ===\n"
    for i, key_point in enumerate(analysis.key_points, start=1):
        text_output += f"{i}. {key_point.point}\n"

    return text_output

# Function to create a PDF report from the analysis
def create_pdf_report(analysis):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Helvetica', '', 12)
    pdf.cell(200, 10, txt="PDF Analysis Report", ln=True, align='C')
    pdf.cell(200, 10, txt=f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True, align='C')
    
    # Convert JSON to text
    clean_text = json_to_text(analysis)
    pdf.multi_cell(0, 10, txt=clean_text)
    return pdf.output(dest='S')  # Return PDF as bytes

# Function to create a Word report from the analysis
def create_word_report(analysis):
    doc = Document()
    doc.add_heading('PDF Analysis Report', 0)
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # Convert JSON to text
    clean_text = json_to_text(analysis)
    doc.add_heading('Analysis', level=1)
    doc.add_paragraph(clean_text)
    
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes.getvalue()

# Streamlit app configuration
st.set_page_config(page_title="Kanoon ki Pehchaan", page_icon="⚖️")

# Custom CSS for styling
def local_css():
    st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@400;700&family=Orbitron:wght@400;700&display=swap');
        body {
            font-family: 'Montserrat', sans-serif;
            background: linear-gradient(135deg, #0A0A0A 0%, #1A1A1A 100%);
            color: #FFFFFF;
        }
        h1, h2, h3 { font-family: 'Orbitron', sans-serif; }
        .main-header {
            position: fixed;
            top: 0;
            width: 100%;
            text-align: center;
            padding: 1.5rem;
            background: rgba(0, 0, 0, 0.8);
            border-bottom: 2px solid #00FFFF;
            box-shadow: 0 0 15px #00FFFF;
            animation: slideInLeft 0.5s ease-in;
            z-index: 1000;
        }
        .flag-stripe {
            height: 6px;
            background: linear-gradient(90deg, #FF00FF 33%, #00FFFF 66%, #00FF00 100%);
            animation: slideInLeft 0.5s ease-in;
        }
        .stTextInput > div > input {
            border-radius: 20px;
            padding: 0.8rem 2rem;
            background: rgba(0, 0, 0, 0.7);
            border: 2px solid #00FFFF;
            color: #FFFFFF;
            transition: all 0.3s ease;
        }
        .stTextInput > div > input:focus {
            border-color: #FF00FF;
            box-shadow: 0 0 15px #FF00FF;
        }
        .stButton > button {
            border-radius: 20px;
            padding: 0.6rem 1.5rem;
            background: linear-gradient(135deg, #00FFFF, #FF00FF);
            color: #000000;
            border: none;
            font-weight: bold;
            text-transform: uppercase;
            box-shadow: 0 0 10px #00FFFF;
            transition: all 0.3s ease;
        }
        .stButton > button:hover {
            transform: scale(1.05);
            box-shadow: 0 0 20px #FF00FF;
        }
        .card {
            background: rgba(255, 255, 255, 0.1);
            backdrop-filter: blur(10px);
            border-radius: 15px;
            border: 1px solid rgba(0, 255, 255, 0.3);
            padding: 1.5rem;
            margin: 1rem 0;
            transition: all 0.3s ease;
        }
        .card:hover {
            transform: translateY(-5px);
            box-shadow: 0 0 20px #FF00FF;
        }
        .footer {
            position: fixed;
            bottom: 0;
            width: 100%;
            background: rgba(0, 0, 0, 0.9);
            padding: 1rem;
            text-align: center;
            border-top: 2px solid #00FFFF;
            animation: fadeIn 0.5s ease-in;
        }
        @keyframes fadeIn { from { opacity: 0; } to { opacity: 1; } }
        @keyframes slideInLeft { from { transform: translateX(-100%); } to { transform: translateX(0); } }
    </style>
    """, unsafe_allow_html=True)

local_css()

# Initialize session state variables
if "current_file" not in st.session_state:
    st.session_state.current_file = None
if "pdf_summary" not in st.session_state:
    st.session_state.pdf_summary = None
if "analysis_time" not in st.session_state:
    st.session_state.analysis_time = 0
if "pdf_report" not in st.session_state:
    st.session_state.pdf_report = None
if "word_report" not in st.session_state:
    st.session_state.word_report = None

# Header and title
st.markdown('<div class="main-header">', unsafe_allow_html=True)
st.markdown('<div class="flag-stripe"></div>', unsafe_allow_html=True)
st.title("⚖️ Kanoon ki Pehchaan")
st.caption("Your AI-powered Legal Document Analyzer")
st.markdown('</div>', unsafe_allow_html=True)

# File uploader and analysis logic
st.markdown('<div class="card animate-fadeIn">', unsafe_allow_html=True)
uploaded_file = st.file_uploader("Upload a PDF file", type="pdf")
if uploaded_file is not None:
    if st.session_state.current_file != uploaded_file.name:
        st.session_state.current_file = uploaded_file.name
        st.session_state.pdf_summary = None
        st.session_state.pdf_report = None
        st.session_state.word_report = None
    text = extract_text_from_pdf(uploaded_file)
    if st.button("Analyze Text"):
        start_time = time.time()
        with st.spinner("Analyzing..."):
            analysis = analyze_text_structured(text)
            st.session_state.pdf_summary = analysis
            st.session_state.pdf_report = create_pdf_report(analysis)
            st.session_state.word_report = create_word_report(analysis)
        end_time = time.time()
        st.session_state.analysis_time = end_time - start_time
        st.subheader("Analysis Results")
        
        # Display analysis in text format
        st.text(json_to_text(analysis))
        
        st.download_button(
            label="Download PDF Report",
            data=st.session_state.pdf_report,
            file_name="analysis_report.pdf",
            mime="application/pdf"
        )
        st.download_button(
            label="Download Word Report",
            data=st.session_state.word_report,
            file_name="analysis_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
st.markdown('</div>', unsafe_allow_html=True)

# Footer with analysis time
st.markdown(f'<div class="footer">Analysis Time: {st.session_state.analysis_time:.1f}s | Powered by Google Generative AI</div>', unsafe_allow_html=True)